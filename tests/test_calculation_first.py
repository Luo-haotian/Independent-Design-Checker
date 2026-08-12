from __future__ import annotations

import json
import zipfile
from pathlib import Path

import pymupdf as fitz
from docx import Document

from idc.artifacts import STANDARD_FILES, create_standard_package
from idc.code_library import library_inventory, lookup_clause
from idc.comments import parse_model_review, validate_comment_codes
from idc.ingestion import ingest_pdf
from idc.llm_review import review_document
from idc.models import (
    CodeBasis,
    CommentAssessment,
    NormalizedSubmission,
    PageClassification,
    PageRole,
    ReviewComment,
    ReviewRun,
    SourceEvidence,
)
from idc.submission import normalize_submission
from report_generator import generate_report_docx


def _pdf(path: Path, pages: list[tuple[str, bool]]) -> Path:
    document = fitz.open()
    for text, landscape in pages:
        page = document.new_page(width=842 if landscape else 595, height=595 if landscape else 842)
        if text:
            page.insert_textbox(fitz.Rect(40, 40, page.rect.width - 40, page.rect.height - 40), text, fontsize=10)
    document.save(path)
    document.close()
    return path


def test_calculation_first_page_classification_matches_submission_shape(tmp_path):
    path = _pdf(
        tmp_path / "submission.pdf",
        [
            ("TEMPORARY WORKS SUBMITTAL FORM Submitted by Contractor Target Approval", False),
            ("Design Calculation Sheet Calc. No. 1 Reference Codes and Standards", False),
            ("Design Calculation Sheet Design Loading bending moment shear force", False),
            ("Typical section for falsework platform", False),
            ("Design Calculation Sheet Member Check deflection utilisation", False),
            ("Design Calculation Sheet Check beam design loading bending moment", False),
            ("Design Calculation Sheet Check joist shear force deflection", False),
            ("Design Calculation Sheet Check bearer bending moment utilisation", False),
            ("Design Calculation Sheet Check scaffold axial load", False),
            ("DRAWING NO. TW-001 LAYOUT PLAN SCALE 1:200 REVISIONS KEY PLAN", True),
            ("rebar weight reference material data sheet", False),
            ("rebar weight reference 3 bundle = 60 kN", False),
        ],
    )
    normalized = normalize_submission(ingest_pdf(path), "temporary")
    assert normalized.calculation_pages == list(range(2, 10))
    assert normalized.drawing_pages == [10]
    assert normalized.supporting_pages == [11, 12]
    assert 10 not in normalized.reviewed_pages


def test_toc_ranges_support_non_contiguous_sections(tmp_path):
    path = _pdf(
        tmp_path / "toc.pdf",
        [
            ("TABLE OF CONTENTS\nDesign Calculations ........ 2-3\nDrawings ........ 4-4\nAppendix ........ 5-6", False),
            ("Member values", False),
            ("Member values", False),
            ("Plan", True),
            ("Data", False),
            ("Data", False),
        ],
    )
    normalized = normalize_submission(ingest_pdf(path), "building")
    assert normalized.pages[0].role == PageRole.TOC
    assert normalized.calculation_pages == [2, 3]
    assert normalized.drawing_pages == [4]
    assert normalized.supporting_pages == [5, 6]


def test_local_code_indexes_and_missing_edition_gate():
    inventory = library_inventory()
    assert {item["family"] for item in inventory if item["available"]} == {"concrete", "foundation", "steel"}
    evidence = lookup_clause("HK Code of Practice for Structural Use of Concrete 2013", "Clause 6.1.2", edition_confirmed=False)
    assert evidence.verified is True
    assert evidence.printed_page == 45
    assert evidence.edition_confirmed is False

    parsed = parse_model_review(
        json.dumps(
            {
                "executive_summary": "Beam calculation reviewed.",
                "comments": [
                    {
                        "location": "PDF page 5, Beam B1",
                        "pages": [5],
                        "submitted_content": "Beam formula",
                        "basis_and_comment": "Candidate check.",
                        "required_action": "Confirm.",
                        "assessment": "ACCEPTABLE",
                        "confidence": 0.9,
                        "code_name": "Concrete Code 2013",
                        "code_clause": "6.1.2",
                    }
                ],
            }
        ),
        [5],
    )
    declarations = [
        {
            "canonical_name": "HK Code of Practice for Structural Use of Concrete 2013 (2020 Edition)",
            "reported_text": "Concrete Code 2013",
            "family": "concrete",
        }
    ]
    comment = validate_comment_codes(parsed.comments, declarations)[0]
    assert comment.code_evidence[0].verified is True
    assert comment.assessment == CommentAssessment.PENDING_CONFIRMATION


def test_missing_or_external_code_never_becomes_acceptable():
    parsed = parse_model_review(
        json.dumps(
            {
                "comments": [
                    {
                        "location": "PDF page 2",
                        "pages": [2],
                        "submitted_content": "Formula",
                        "basis_and_comment": "Appears adequate.",
                        "required_action": "None.",
                        "assessment": "ACCEPTABLE",
                        "confidence": 0.8,
                    },
                    {
                        "location": "PDF page 3",
                        "pages": [3],
                        "submitted_content": "Timber check",
                        "basis_and_comment": "Candidate BS check.",
                        "required_action": "Confirm.",
                        "assessment": "ACCEPTABLE",
                        "confidence": 0.8,
                        "code_name": "BS 5268-2:2002",
                        "code_clause": "2.3",
                    },
                ]
            }
        ),
        [2, 3],
    )
    declarations = [{"canonical_name": "BS 5268-2 : 2002", "reported_text": "BS 5268-2 : 2002", "family": "unclassified"}]
    comments = validate_comment_codes(parsed.comments, declarations)
    assert [item.assessment for item in comments] == [
        CommentAssessment.INFORMATION_REQUIRED,
        CommentAssessment.INFORMATION_REQUIRED,
    ]
    assert comments[1].code_evidence[0].verified is False


def test_chunks_merge_into_one_comment_set_without_duplicate_reports(tmp_path):
    path = _pdf(
        tmp_path / "long.pdf",
        [("Design Calculation Sheet " + ("member check bending moment " * 70), False) for _ in range(3)],
    )
    extraction = ingest_pdf(path)
    calls = 0

    def model(_prompt: str) -> str:
        nonlocal calls
        calls += 1
        return json.dumps(
            {
                "executive_summary": "Calculation pages reviewed.",
                "comments": [
                    {
                        "location": "Calculation member",
                        "pages": [calls],
                        "submitted_content": f"Issue {calls}",
                        "basis_and_comment": "Evidence-based calculation issue.",
                        "required_action": "Provide calculation evidence.",
                        "assessment": "INFORMATION_REQUIRED",
                        "confidence": 0.75,
                    }
                ],
            }
        )

    summary, raw, covered, missing, executive, comments = review_document(
        extraction,
        "Review calculation content and return JSON.\n{content}",
        model,
        max_input_chars=1200,
    )
    assert calls > 1
    assert covered == [1, 2, 3]
    assert missing == []
    assert len(raw) == calls
    assert len(comments) == calls
    assert summary.count("# Calculation Review Internal Summary") == 1
    assert "Review chunk" not in summary
    assert executive == "Calculation pages reviewed."


def _review_run(pdf: Path) -> ReviewRun:
    structure = NormalizedSubmission(
        review_profile="temporary",
        pages=[
            PageClassification(1, PageRole.CALCULATION, "Calculation", 0.95, ["calculation heading"]),
            PageClassification(2, PageRole.DRAWING, "Drawing", 0.95, ["drawing title block"]),
        ],
        calculation_pages=[1],
        drawing_pages=[2],
        reviewed_pages=[1],
        deferred_pages=[2],
    )
    return ReviewRun(
        run_id="run1",
        source_file=str(pdf),
        source_sha256="secret-hash-not-for-human-report",
        code_basis=CodeBasis(unresolved_codes=["internal code-basis issue"]),
        submission_structure=structure,
        executive_summary="Calculation review identified one item requiring clarification.",
        comments=[
            ReviewComment(
                comment_no=1,
                location="PDF page 1, Beam B1",
                submitted_content="The design formula has no stated code clause.",
                basis_and_comment="No applicable design code was identified.",
                required_action="State the applicable code, edition and clause.",
                assessment=CommentAssessment.INFORMATION_REQUIRED,
                confidence=0.72,
                evidence=[SourceEvidence(str(pdf), page=1, confidence=0.72)],
            )
        ],
        model_provider="secret-provider",
        model_name="secret-model",
    )


def test_standard_package_and_human_word_report_are_separated(tmp_path):
    pdf = _pdf(tmp_path / "input.pdf", [("Design Calculation Sheet", False), ("DRAWING NO. A1", True)])
    extraction = ingest_pdf(pdf)
    run = _review_run(pdf)
    package = create_standard_package(run, extraction, tmp_path / "standard.zip")
    with zipfile.ZipFile(package) as archive:
        assert set(archive.namelist()) == set(STANDARD_FILES)
        processing = archive.read("processing_record.json").decode("utf-8")
        assert "secret-hash-not-for-human-report" in processing
        assert "input.pdf" not in archive.namelist()

    docx = tmp_path / "report.docx"
    generate_report_docx(
        "internal maintenance content",
        str(pdf),
        "temporary",
        str(docx),
        "secret-model",
        "secret-provider",
        project_title="Sanitized Project",
        checked_item="Sanitized Calculation",
        job_reference="TEST-001",
        review_run=run,
    )
    document = Document(docx)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "IDC Review Comments" in text
    assert "Required action / assessment / confidence" in text
    for forbidden in (
        "Code basis:",
        "Basis profile:",
        "Deterministic rule pack:",
        "Source SHA-256:",
        "secret-model",
        "secret-provider",
        "secret-hash-not-for-human-report",
    ):
        assert forbidden not in text
