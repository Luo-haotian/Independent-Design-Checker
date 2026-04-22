"""Professional Word report generator for IDC."""

import os
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)")
NUMBERED_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*\.?)\s+(.+)$")
INLINE_BOLD_RE = re.compile(r"(\*\*.+?\*\*)")
MEINHARDT_BLUE = "154D93"
MEINHARDT_RED = "F01E14"
ACCENT_BLUE = "2E37FF"
LIGHT_GREY = RGBColor(153, 153, 153)
RED_PLACEHOLDER = RGBColor(192, 0, 0)


def _asset_path(name: str) -> Path | None:
    """Resolve a bundled or local report asset path."""
    candidates = []
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        candidates.append(Path(sys._MEIPASS) / "report_assets" / name)
    base_dir = Path(__file__).resolve().parent
    candidates.append(base_dir / "report_assets" / name)
    candidates.append(base_dir / "report sample" / name)

    for path in candidates:
        if path.exists():
            return path
    return None


def _set_run_font(
    run,
    size: float,
    *,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    font_name: str = "Arial",
):
    """Apply a consistent font setup."""
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color is not None:
        run.font.color.rgb = color


def _set_cell_shading(cell, fill: str):
    """Apply background shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_borders(cell, color: str = "000000", size: int = 8):
    """Apply simple borders to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _clear_cell_borders(cell):
    """Remove visible borders from a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_table_width(table, width_cm: float):
    """Set a fixed total table width."""
    table.autofit = False
    for column in table.columns:
        column.width = Cm(width_cm / len(table.columns))


def _set_page_layout(section):
    """Apply A4 layout and margins."""
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)


def _unlink_section_headers_footers(section):
    """Break header/footer linkage so each section can be styled independently."""
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False


def _clean_text(value: str | None) -> str:
    """Normalize whitespace for display."""
    if not value:
        return ""
    return re.sub(r"\s+", " ", value).strip()


def _fallback_title_from_pdf(pdf_path: str) -> str:
    """Build a readable fallback title from the file name."""
    stem = os.path.splitext(os.path.basename(pdf_path))[0]
    title = re.sub(r"[_\-]+", " ", stem)
    title = re.sub(r"\s+", " ", title).strip()
    return title or "Structural Submission"


def _wrap_lines(text: str, max_chars: int, *, uppercase: bool = False) -> list[str]:
    """Wrap text into display lines."""
    source = text.upper() if uppercase else text
    words = source.split()
    if not words:
        return []

    lines = []
    current = []
    for word in words:
        test = " ".join(current + [word])
        if current and len(test) > max_chars:
            lines.append(" ".join(current))
            current = [word]
        else:
            current.append(word)
    if current:
        lines.append(" ".join(current))
    return lines


def _wrap_upper_lines(text: str, max_chars: int) -> list[str]:
    """Wrap text into uppercase display lines."""
    return _wrap_lines(text, max_chars, uppercase=True)


def _display_subject(checked_item: str | None, project_title: str | None, pdf_path: str) -> str:
    """Choose the most useful subject text for the report title."""
    subject = _clean_text(checked_item) or _clean_text(project_title) or _fallback_title_from_pdf(pdf_path)
    subject = re.sub(r"^(design of|checking of)\s+", "", subject, flags=re.IGNORECASE)
    subject = re.sub(r"^the\s+", "", subject, flags=re.IGNORECASE)
    return subject


def _cover_title_lines(subject: str) -> list[str]:
    """Build a natural IDC cover title."""
    lines = [
        "INDEPENDENT DESIGN",
        "CHECKING REPORT",
        "",
        "FOR",
    ]
    lines.extend(_wrap_upper_lines(subject, 30)[:4])
    return lines


def _initial_counters() -> dict[int, int]:
    """Return empty heading counters."""
    return {1: 0, 2: 0, 3: 0}


def _apply_number_to_counters(number: str, counters: dict[int, int]) -> tuple[int, str]:
    """Update counters from an explicit numbered heading."""
    parts = [int(part) for part in number.rstrip(".").split(".") if part]
    level = min(len(parts), 3)
    for idx in range(3):
        counters[idx + 1] = parts[idx] if idx < len(parts) else 0
    normalized = ".".join(str(part) for part in parts[:level])
    return level, normalized


def _next_number(level: int, counters: dict[int, int]) -> str:
    """Advance counters and return the new heading number."""
    counters[level] += 1
    for deeper in range(level + 1, 4):
        counters[deeper] = 0
    return ".".join(str(counters[idx]) for idx in range(1, level + 1))


def _line_cost(text: str) -> int:
    """Approximate how many body lines a paragraph will consume."""
    return max(1, (len(text) // 75) + 1)


def _normalize_content(content: str) -> tuple[str, list[dict[str, object]]]:
    """Normalize headings, add numbering, and estimate TOC page numbers."""
    lines = content.splitlines()
    counters = _initial_counters()
    normalized_lines: list[str] = []
    toc_entries: list[dict[str, object]] = []
    current_page = 4
    used_lines = 0
    page_capacity = 34

    for raw_line in lines:
        stripped = raw_line.strip()
        if not stripped:
            normalized_lines.append("")
            continue

        markdown_match = MARKDOWN_HEADING_RE.match(stripped)
        if markdown_match:
            level = min(len(markdown_match.group(1)), 3)
            title = markdown_match.group(2).strip()
            if not toc_entries and level == 1 and "report" in title.lower():
                continue
            if not toc_entries and counters[1] == 0 and level > 1:
                level = 1
            explicit_match = NUMBERED_HEADING_RE.match(title)
            if explicit_match:
                level, number = _apply_number_to_counters(explicit_match.group(1), counters)
                clean_title = explicit_match.group(2).strip()
            else:
                number = _next_number(level, counters)
                clean_title = title

            display = f"{number}. {clean_title}" if level == 1 else f"{number} {clean_title}"
            toc_entries.append(
                {
                    "level": level,
                    "number": number,
                    "title": clean_title,
                    "display": display,
                    "page": current_page,
                }
            )
            normalized_lines.append(f"{'#' * level} {display}")
            used_lines += 2
        else:
            normalized_lines.append(stripped)
            used_lines += _line_cost(stripped)

        while used_lines > page_capacity:
            used_lines -= page_capacity
            current_page += 1

    return "\n".join(normalized_lines), toc_entries


def _add_logo(paragraph, *, width_cm: float = 3.8, asset_name: str = "header_logo_clean.png"):
    """Add the Meinhardt logo."""
    asset = _asset_path(asset_name)
    if asset:
        run = paragraph.add_run()
        run.add_picture(str(asset), width=Cm(width_cm))
        return

    parts = [("ME", RGBColor(21, 77, 147)), ("I", RGBColor(240, 30, 20)), ("NHARDT", RGBColor(21, 77, 147))]
    for text, color in parts:
        run = paragraph.add_run(text)
        _set_run_font(run, 22, bold=True, color=color)


def _add_cover_page(doc: Document, subject: str, construction_title: str):
    """Add the sample-inspired cover page."""
    usable_width = 18.2
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(cover_table, usable_width)
    cell = cover_table.cell(0, 0)
    _set_cell_shading(cell, MEINHARDT_BLUE)
    _clear_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cover_table.rows[0].height = Cm(14.8)
    cover_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    logo_p = cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_p.paragraph_format.space_before = Pt(22)
    logo_p.paragraph_format.space_after = Pt(38)
    _add_logo(logo_p, width_cm=3.5, asset_name="meinhardt_logo.png")

    cover_lines = _cover_title_lines(subject)
    for index, line in enumerate(cover_lines):
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Cm(0.1)
        paragraph.paragraph_format.space_after = Pt(5 if line else 10)
        if line:
            _set_run_font(
                paragraph.add_run(line),
                18 if index < 2 else 16,
                bold=True,
                color=RGBColor(255, 255, 255),
            )

    if construction_title:
        label = cell.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label.paragraph_format.left_indent = Cm(0.1)
        label.paragraph_format.space_before = Pt(8)
        label.paragraph_format.space_after = Pt(3)
        _set_run_font(label.add_run("CONSTRUCTION TITLE"), 10.5, bold=True, color=RGBColor(255, 255, 255))

        for line in _wrap_upper_lines(construction_title, 34)[:3]:
            paragraph = cell.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Cm(0.1)
            paragraph.paragraph_format.space_after = Pt(2)
            _set_run_font(paragraph.add_run(line), 11.5, color=RGBColor(255, 255, 255))

    section_break = cell.add_paragraph()
    section_break.paragraph_format.space_after = Pt(10)

    for index, line in enumerate(["(PART 1)", "(VOL 1 OF 1)"]):
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Cm(0.1)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.space_before = Pt(0 if index else 2)
        _set_run_font(
            paragraph.add_run(line),
            15,
            bold=True,
            color=RGBColor(255, 255, 255),
        )

    date_paragraph = cell.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_paragraph.paragraph_format.left_indent = Cm(0.1)
    date_paragraph.paragraph_format.space_before = Pt(8)
    _set_run_font(
        date_paragraph.add_run(datetime.now().strftime("%B %Y").upper()),
        13,
        color=RGBColor(255, 255, 255),
    )

    stripe = doc.add_table(rows=1, cols=1)
    stripe.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(stripe, usable_width)
    stripe_cell = stripe.cell(0, 0)
    _set_cell_shading(stripe_cell, MEINHARDT_RED)
    _clear_cell_borders(stripe_cell)
    stripe.rows[0].height = Cm(0.7)
    stripe.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    spacer = doc.add_table(rows=1, cols=1)
    spacer.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(spacer, usable_width)
    _clear_cell_borders(spacer.cell(0, 0))
    spacer.rows[0].height = Cm(4.8)
    spacer.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    info = doc.add_table(rows=1, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(info, usable_width)
    left = info.cell(0, 0)
    right = info.cell(0, 1)
    for cell in (left, right):
        _clear_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    paragraph = left.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(paragraph.add_run("Submitted to"), 12, bold=True)
    submitted_asset = _asset_path("submitted_to.png")
    if submitted_asset:
        picture = left.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.LEFT
        picture.add_run().add_picture(str(submitted_asset), width=Cm(5.4))
    else:
        paragraph = left.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(paragraph.add_run("Hong Kong Housing Authority"), 10)

    paragraph = right.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(paragraph.add_run("Prepared By"), 12, bold=True)
    paragraph = right.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(paragraph.add_run("Meinhardt (C&S) Ltd"), 10)


def _add_cover_section_break(doc: Document):
    """Insert a new section after the cover."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_page_layout(section)
    _unlink_section_headers_footers(section)
    section.different_first_page_header_footer = True
    return section


def _configure_second_page_header(section):
    """Page 2 header: logo only."""
    header = section.first_page_header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(2)
    _add_logo(paragraph, width_cm=3.9, asset_name="header_logo_clean.png")


def _configure_body_header_footer(section, subject: str, construction_title: str, job_reference: str):
    """Configure body header and footer for pages after the revision page."""
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(18.2))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(13.6)
    table.columns[1].width = Cm(4.6)
    for cell in table.row_cells(0):
        _clear_cell_borders(cell)

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    title_paragraph = left.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.space_after = Pt(1)
    _set_run_font(
        title_paragraph.add_run("Independent Design Checking Report"),
        10,
        bold=True,
        color=LIGHT_GREY,
        font_name="Times New Roman",
    )

    for line in _wrap_lines(subject, 54)[:2]:
        paragraph = left.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(1)
        _set_run_font(
            paragraph.add_run(line),
            9.5,
            bold=True,
            color=LIGHT_GREY,
            font_name="Times New Roman",
        )

    if construction_title:
        for line in _wrap_lines(construction_title, 54)[:2]:
            paragraph = left.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(1)
            _set_run_font(
                paragraph.add_run(line),
                9,
                color=LIGHT_GREY,
                font_name="Times New Roman",
            )

    logo_paragraph = right.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_logo(logo_paragraph, width_cm=3.7, asset_name="header_logo_clean.png")

    line = header.add_paragraph()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_BLUE)
    border.append(bottom)
    line._p.get_or_add_pPr().append(border)

    footer = section.footer
    footer_line = footer.paragraphs[0]
    footer_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    border = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "18")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), ACCENT_BLUE)
    border.append(top)
    footer_line._p.get_or_add_pPr().append(border)

    info = footer.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(info.add_run(f"Job No. {job_reference}\nMeinhardt (C&S) Limited"), 9)


def _add_revision_page(doc: Document, subject: str, construction_title: str, job_reference: str):
    """Add the revision and approval page."""
    for _ in range(2):
        doc.add_paragraph()

    layout = doc.add_table(rows=1, cols=2)
    layout.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(layout, 18.2)
    for cell in layout.row_cells(0):
        _clear_cell_borders(cell)

    right = layout.cell(0, 1)
    title_lines = [
        "INDEPENDENT DESIGN CHECKING REPORT",
        f"FOR {subject.upper()}",
        "",
        "",
        "(PART 1)",
        "(VOL 1 OF 1)",
    ]
    for idx, line in enumerate(title_lines):
        paragraph = right.paragraphs[0] if idx == 0 else right.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(3 if line else 8)
        if line:
            _set_run_font(
                paragraph.add_run(line),
                14 if idx < 2 else 12.5,
                bold=True,
            )

    if construction_title:
        for line in _wrap_upper_lines(construction_title, 34)[:3]:
            paragraph = right.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(2)
            _set_run_font(paragraph.add_run(line), 11.5)

    for _ in range(2):
        doc.add_paragraph()

    table = doc.add_table(rows=9, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Date", "Revision", "Prepared", "", "Checked", "", "Approved", ""]
    widths = [2.0, 2.0, 1.9, 2.0, 1.9, 2.1, 1.9, 2.1]

    for idx, width in enumerate(widths):
        table.columns[idx].width = Cm(width)

    for row_idx, row in enumerate(table.rows):
        row.height = Cm(0.82)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for col_idx, cell in enumerate(row.cells):
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if row_idx == 0:
                _set_run_font(paragraph.add_run(headers[col_idx]), 10.5, bold=True)

    first_row = [
        datetime.now().strftime("%m/%Y"),
        "First\nIssue",
        "XX",
        "TO BE\nUPDATED",
        "XX",
        "TO BE\nUPDATED",
        "XX",
        "TO BE\nUPDATED",
    ]
    revisions = ["OA", "OB", "OC", "OD", "OE", "OF", "OG"]

    for col_idx, value in enumerate(first_row):
        paragraph = table.cell(1, col_idx).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        color = RED_PLACEHOLDER if value == "XX" or "UPDATED" in value else None
        _set_run_font(paragraph.add_run(value), 10, color=color)

    for row_idx, code in enumerate(revisions, start=2):
        paragraph = table.cell(row_idx, 1).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.add_run(code), 10)

    project_paragraph = doc.add_paragraph()
    project_paragraph.paragraph_format.space_before = Pt(8)
    _set_run_font(project_paragraph.add_run(f"Project / Job No.: {job_reference}"), 11)


def _add_toc_page(doc: Document, toc_entries: list[dict[str, object]]):
    """Add a TOC page with hierarchical numbering and page numbers."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("TABLE OF CONTENTS")
    _set_run_font(run, 20)
    run.underline = True

    if not toc_entries:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.add_run("(Report content will appear after this page)"), 11, italic=True)
        return

    table = doc.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(2.0)
    table.columns[1].width = Cm(13.5)
    table.columns[2].width = Cm(1.6)

    for entry in toc_entries:
        row = table.add_row().cells
        for cell in row:
            _clear_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        level = int(entry["level"])
        number = str(entry["number"])
        title_text = str(entry["title"])
        page = str(entry["page"])

        number_paragraph = row[0].paragraphs[0]
        title_paragraph = row[1].paragraphs[0]
        page_paragraph = row[2].paragraphs[0]

        title_paragraph.paragraph_format.left_indent = Cm(0.7 * (level - 1))
        title_paragraph.paragraph_format.space_after = Pt(2)
        page_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if level == 1:
            _set_run_font(number_paragraph.add_run(f"{number}."), 15, bold=True)
            _set_run_font(title_paragraph.add_run(title_text.upper()), 15, bold=True)
            _set_run_font(page_paragraph.add_run(page), 14, bold=True)
        else:
            _set_run_font(number_paragraph.add_run(number), 11)
            _set_run_font(title_paragraph.add_run(title_text), 11)
            _set_run_font(page_paragraph.add_run(page), 11)


def _apply_heading_style(doc: Document, paragraph, level: int, text: str):
    """Apply heading styling with Word built-in heading styles."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    style_name = f"Heading {min(level, 3)}"
    try:
        paragraph.style = doc.styles[style_name]
    except KeyError:
        pass

    size = 16 if level == 1 else 13 if level == 2 else 11
    _set_run_font(paragraph.add_run(text), size, bold=True)


def _add_inline_formatted_runs(
    paragraph,
    text: str,
    size: float,
    *,
    color: RGBColor | None = None,
    font_name: str = "Arial",
):
    """Render simple Markdown bold markers as Word bold runs."""
    parts = INLINE_BOLD_RE.split(text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**") and len(part) > 4
        run_text = part[2:-2] if is_bold else part
        _set_run_font(paragraph.add_run(run_text), size, bold=is_bold, color=color, font_name=font_name)


def _add_normal_paragraph(doc: Document, text: str):
    """Add a body paragraph."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_inline_formatted_runs(paragraph, text, 11)


def _add_bullet_paragraph(doc: Document, text: str):
    """Add a bullet paragraph."""
    try:
        paragraph = doc.add_paragraph(style="List Bullet")
    except KeyError:
        paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _add_inline_formatted_runs(paragraph, text, 11)


def _add_numbered_paragraph(doc: Document, text: str):
    """Add a numbered paragraph."""
    try:
        paragraph = doc.add_paragraph(style="List Number")
    except KeyError:
        paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _add_inline_formatted_runs(paragraph, text, 11)


def _parse_and_add_content(doc: Document, content: str):
    """Parse normalized Markdown-like content and add it to the document."""
    lines = content.splitlines()
    in_code_block = False
    code_lines: list[str] = []

    for raw_line in lines:
        stripped = raw_line.strip()
        if stripped.startswith("```"):
            if in_code_block and code_lines:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(6)
                run = paragraph.add_run("\n".join(code_lines))
                _set_run_font(run, 10, font_name="Courier New")
                code_lines = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(raw_line)
            continue

        if not stripped:
            continue

        markdown_match = MARKDOWN_HEADING_RE.match(stripped)
        if markdown_match:
            level = min(len(markdown_match.group(1)), 3)
            _apply_heading_style(doc, doc.add_paragraph(), level, markdown_match.group(2).strip())
            continue

        bullet_match = re.match(r"^([\-\*])\s+(.*)", stripped)
        if bullet_match:
            _add_bullet_paragraph(doc, bullet_match.group(2))
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if numbered_match:
            _add_numbered_paragraph(doc, numbered_match.group(1))
            continue

        _add_normal_paragraph(doc, stripped)


def generate_report_docx(
    content: str,
    pdf_path: str,
    struct_type: str,
    output_path: str,
    model: str,
    provider: str,
    used_ocr: bool = False,
    *,
    project_title: str | None = None,
    checked_item: str | None = None,
    job_reference: str | None = None,
) -> str:
    """Generate a professional .docx report and return the output path."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    subject = _display_subject(checked_item, project_title, pdf_path)
    project_title = _clean_text(project_title) or subject
    job_reference = _clean_text(job_reference) or _fallback_title_from_pdf(pdf_path)[:30].upper()
    normalized_content, toc_entries = _normalize_content(content)

    first_section = doc.sections[0]
    _set_page_layout(first_section)
    _unlink_section_headers_footers(first_section)
    first_section.different_first_page_header_footer = True
    _add_cover_page(doc, subject, project_title)

    second_section = _add_cover_section_break(doc)
    _configure_second_page_header(second_section)
    _add_revision_page(doc, subject, project_title, job_reference)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_page_layout(body_section)
    _unlink_section_headers_footers(body_section)
    _configure_body_header_footer(body_section, subject, project_title, job_reference)
    _add_toc_page(doc, toc_entries)

    doc.add_page_break()
    _parse_and_add_content(doc, normalized_content)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(18)
    run = footer.add_run(
        f"Generated by IDC on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"Provider: {provider.upper()} | Model: {model}"
    )
    _set_run_font(run, 9, italic=True, color=LIGHT_GREY)

    if used_ocr:
        run.add_break(WD_BREAK.LINE)
        ocr_run = footer.add_run("OCR-assisted extraction used for this submission.")
        _set_run_font(ocr_run, 9, italic=True, color=LIGHT_GREY)

    doc.save(output_path)
    return output_path
